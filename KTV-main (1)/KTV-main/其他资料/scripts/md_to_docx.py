"""
Markdown to Word Converter
将Markdown文件转换为Word文档格式
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def create_document():
    """创建Word文档并设置默认样式"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)  # 小四号
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    return doc


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    # 设置标题字体
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 添加表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # 设置表头样式
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '宋体'
                run.font.size = Pt(11)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(11)
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return table


def add_code_block(doc, code):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    # 设置背景色（灰色）
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    run.element.rPr.append(shading)
    return para


def add_bullet_list(doc, items):
    """添加无序列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_numbered_list(doc, items):
    """添加有序列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Number')
        for run in para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def parse_markdown_table(lines, start_idx):
    """解析Markdown表格"""
    headers = []
    rows = []
    idx = start_idx

    # 解析表头
    if idx < len(lines) and '|' in lines[idx]:
        header_line = lines[idx].strip()
        headers = [h.strip() for h in header_line.split('|') if h.strip()]
        idx += 1

    # 跳过分隔线
    if idx < len(lines) and '---' in lines[idx]:
        idx += 1

    # 解析数据行
    while idx < len(lines) and '|' in lines[idx] and '---' not in lines[idx]:
        row_line = lines[idx].strip()
        if row_line:
            row = [r.strip() for r in row_line.split('|') if r.strip()]
            if row:
                rows.append(row)
        idx += 1

    return headers, rows, idx


def convert_md_to_docx(md_file, docx_file):
    """将Markdown转换为Word"""
    doc = create_document()

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    idx = 0
    in_code_block = False
    code_content = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # 处理代码块
        if stripped.startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, '\n'.join(code_content))
                code_content = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_content.append(line)
            idx += 1
            continue

        # 处理空行
        if not stripped:
            idx += 1
            continue

        # 处理标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            if level <= 4:
                add_heading(doc, text, level)
            idx += 1
            continue

        # 处理表格
        if '|' in stripped and idx + 1 < len(lines) and '---' in lines[idx + 1]:
            headers, rows, new_idx = parse_markdown_table(lines, idx)
            if headers and rows:
                add_table(doc, headers, rows)
            idx = new_idx
            continue

        # 处理无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            items = []
            while idx < len(lines) and (lines[idx].strip().startswith('- ') or lines[idx].strip().startswith('* ')):
                item_text = lines[idx].strip()[2:].strip()
                # 处理加粗文本
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                items.append(item_text)
                idx += 1
            add_bullet_list(doc, items)
            continue

        # 处理有序列表
        if re.match(r'^\d+\.', stripped):
            items = []
            while idx < len(lines) and re.match(r'^\d+\.', lines[idx].strip()):
                item_text = re.sub(r'^\d+\.\s*', '', lines[idx].strip())
                # 处理加粗文本
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                items.append(item_text)
                idx += 1
            add_numbered_list(doc, items)
            continue

        # 处理水平线
        if stripped in ['---', '***', '___']:
            doc.add_paragraph('─' * 50)
            idx += 1
            continue

        # 处理普通段落
        text = stripped
        # 处理加粗文本
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # 处理行内代码
        text = re.sub(r'`(.*?)`', r'\1', text)

        if text:
            add_paragraph(doc, text)

        idx += 1

    # 保存文档
    doc.save(docx_file)
    print(f"转换完成！文件已保存到：{docx_file}")


if __name__ == '__main__':
    import sys

    # 默认转换系统操作说明书
    if len(sys.argv) < 3:
        md_file = r'e:\syyktv\KTV-main (1)\KTV-main\docs\系统操作说明书.md'
        docx_file = r'e:\syyktv\KTV-main (1)\KTV-main\docs\系统操作说明书.docx'
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]

    convert_md_to_docx(md_file, docx_file)
