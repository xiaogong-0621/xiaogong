"""
声域友项目说明书 Markdown → Word 转换脚本
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def create_document():
    """创建Word文档并设置默认样式"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置段落间距
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    return doc


def add_heading_styled(doc, text, level):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(22)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
        elif level == 4:
            run.font.size = Pt(13)
    return heading


def add_paragraph_styled(doc, text, bold=False, indent=False):
    """添加带样式的段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if indent:
        para.paragraph_format.first_line_indent = Pt(24)
    return para


def add_table_styled(doc, headers, rows):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = '宋体'
        run.font.size = Pt(11)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = '宋体'
            run.font.size = Pt(11)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()  # 表后空行
    return table


def add_code_block_styled(doc, code):
    """添加代码块"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    run.element.rPr.append(shading)
    return para


def convert_md_to_docx(md_file, docx_file):
    """主转换函数"""
    doc = create_document()

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    idx = 0
    in_code_block = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # 处理代码块
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block_styled(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            idx += 1
            continue

        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue

        # 跳过空行
        if not stripped:
            idx += 1
            continue

        # 水平线
        if stripped.startswith('---'):
            para = doc.add_paragraph()
            run = para.add_run('─' * 40)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)
            idx += 1
            continue

        # 标题
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            if level <= 4:
                add_heading_styled(doc, text, level)
            idx += 1
            continue

        # 表格
        if '|' in stripped and idx + 1 < len(lines) and '---' in lines[idx + 1]:
            # 解析表格
            headers_line = lines[idx].strip()
            headers = [h.strip() for h in headers_line.split('|') if h.strip()]
            idx += 2  # 跳过表头行和分隔行
            rows = []
            while idx < len(lines) and '|' in lines[idx] and lines[idx].strip():
                row_line = lines[idx].strip()
                row = [r.strip() for r in row_line.split('|') if r.strip()]
                if row:
                    rows.append(row)
                idx += 1
            if headers and rows:
                add_table_styled(doc, headers, rows)
            continue

        # 无序列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            items = []
            while idx < len(lines) and (lines[idx].strip().startswith('- ') or lines[idx].strip().startswith('* ')):
                item_text = lines[idx].strip()
                # 提取列表项内容（跳过 -  或 *  ）
                item_text = re.sub(r'^[-*]\s+', '', item_text)
                # 去除所有markdown格式标记
                item_text = re.sub(r'\*\*(.*?)\*\*', r'\1', item_text)
                item_text = re.sub(r'`(.*?)`', r'\1', item_text)
                items.append(item_text)
                idx += 1
            for item in items:
                para = doc.add_paragraph(style='List Bullet')
                # 清除默认文本
                para.clear()
                run = para.add_run(item)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            continue

        # 普通段落
        text = stripped
        # 处理粗体
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # 处理行内代码
        text = re.sub(r'`(.*?)`', r'\1', text)
        # 处理链接 [text](url) -> text
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)

        if text:
            add_paragraph_styled(doc, text)

        idx += 1

    # 保存
    doc.save(docx_file)
    print(f"✅ 转换完成！Word文档已保存到：{docx_file}")


if __name__ == '__main__':
    import sys
    if len(sys.argv) < 3:
        md_file = r'e:\syyktv\KTV-main (1)\KTV-main\声域友项目说明书.md'
        docx_file = r'e:\syyktv\KTV-main (1)\KTV-main\声域友项目说明书.docx'
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]

    convert_md_to_docx(md_file, docx_file)
