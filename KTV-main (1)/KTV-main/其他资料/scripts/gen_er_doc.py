"""生成声域友 KTV 系统 ER 图 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("声域友 KTV 线上社交平台")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 100, 200)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("数据库 ER 图设计文档")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("项目版本：v2.0.8\n技术栈：ASP.NET Core 8.0 + Vue 3 + SQL Server")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ===== 目录页 =====
doc.add_heading("目录", level=1)
toc_items = [
    "1. 系统概述",
    "2. 实体列表",
    "3. 实体关系总览",
    "4. 各实体详细字段",
    "5. 关系说明",
    "6. 状态值枚举",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ===== 1. 系统概述 =====
doc.add_heading("1. 系统概述", level=1)
doc.add_paragraph(
    "声域友是一套线上社交 KTV 平台，用户通过浏览器或微信小程序创建或加入虚拟房间，"
    "在房间内与朋友一起点歌、排队、实时聊天，享受线上 KTV 体验。"
    "系统分为 Web 用户端、微信小程序端和管理后台三端，共用同一个后端 API 和数据库。"
)

doc.add_heading("技术架构", level=2)
tech_table = doc.add_table(rows=5, cols=2)
tech_table.style = "Light Grid Accent 1"
headers = [("层", "技术")]
data = [
    ("后端", "ASP.NET Core 8.0, Dapper, SignalR, JWT"),
    ("前端", "Vue 3, TypeScript, Vite 8, Tailwind CSS 3, Element Plus, Pinia"),
    ("小程序", "微信小程序原生框架, WebSocket"),
    ("数据库", "SQL Server (localhost\\SQLEXPRESS)"),
]
for i, (h1, h2) in enumerate(headers + data):
    cells = tech_table.rows[i].cells
    cells[0].text = h1
    cells[1].text = h2
    if i == 0:
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

doc.add_paragraph()

# ===== 2. 实体列表 =====
doc.add_heading("2. 实体列表", level=1)
doc.add_paragraph("系统共包含 11 张数据表：")

entity_table = doc.add_table(rows=12, cols=4)
entity_table.style = "Light Grid Accent 1"
entity_headers = ["序号", "表名", "中文名", "用途"]
entities = [
    ("1", "Users", "用户表", "存储所有用户信息（普通用户和管理员）"),
    ("2", "Songs", "歌曲表", "平台歌曲库，支持 MP3/FLAC 格式"),
    ("3", "Rooms", "虚拟房间表", "用户创建的线上 KTV 房间"),
    ("4", "RoomUsers", "房间用户关联表", "记录当前在房间内的用户"),
    ("5", "RoomRequests", "房间申请表", "用户提交的开房申请"),
    ("6", "PlayQueue", "播放队列表", "房间内的歌曲播放队列"),
    ("7", "Favorites", "收藏表", "用户收藏的歌曲"),
    ("8", "Feedbacks", "反馈表", "用户提交的歌曲反馈"),
    ("9", "ChatMessages", "聊天消息表", "房间内的实时聊天记录"),
    ("10", "SystemSettings", "系统设置表", "平台配置项（键值对）"),
    ("11", "OperationLogs", "操作日志表", "管理员操作记录"),
]

for j, h in enumerate(entity_headers):
    cell = entity_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

for i, row_data in enumerate(entities):
    for j, val in enumerate(row_data):
        entity_table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# ===== 3. 实体关系总览 =====
doc.add_heading("3. 实体关系总览", level=1)
doc.add_paragraph("以下为各实体之间的外键关联关系：")

rel_table = doc.add_table(rows=10, cols=4)
rel_table.style = "Light Grid Accent 1"
rel_headers = ["父表", "子表", "关系类型", "说明"]
relations = [
    ("Users", "Rooms", "一对多", "一个用户可创建多个房间"),
    ("Users", "RoomUsers", "一对多", "一个用户加入房间（UNIQUE 约束：同时只在一个房间）"),
    ("Users", "PlayQueue", "一对多", "一个用户可点播多首歌"),
    ("Users", "Favorites", "一对多", "一个用户可收藏多首歌"),
    ("Users", "ChatMessages", "一对多", "一个用户可发送多条消息"),
    ("Users", "RoomRequests", "一对多", "一个用户可提交多次开房申请"),
    ("Users", "Feedbacks", "一对多", "一个用户可提交多次反馈"),
    ("Rooms", "PlayQueue", "一对多", "一个房间有一个播放队列"),
    ("Songs", "PlayQueue", "一对多", "一首歌可被多个房间点播"),
]

for j, h in enumerate(rel_headers):
    cell = rel_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

for i, row_data in enumerate(relations):
    for j, val in enumerate(row_data):
        rel_table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# ===== 4. 各实体详细字段 =====
doc.add_heading("4. 各实体详细字段", level=1)

table_defs = [
    ("4.1 Users 用户表", [
        ("Id", "INT", "PK", "自增主键"),
        ("Username", "NVARCHAR(50)", "UNIQUE", "用户名"),
        ("PasswordHash", "NVARCHAR(256)", "", "密码哈希"),
        ("DisplayName", "NVARCHAR(100)", "", "显示名称"),
        ("Phone", "NVARCHAR(20)", "", "手机号"),
        ("Email", "NVARCHAR(100)", "UNIQUE", "邮箱"),
        ("AvatarUrl", "NVARCHAR(500)", "", "头像路径"),
        ("Balance", "DECIMAL(10,2)", "", "余额"),
        ("IsVip", "BIT", "", "是否 VIP"),
        ("Role", "NVARCHAR(20)", "", "角色 (user/admin)"),
        ("Status", "NVARCHAR(20)", "", "状态 (active/disabled)"),
        ("LastActiveAt", "DATETIME2", "", "最后活跃时间"),
        ("CreatedAt", "DATETIME2", "", "创建时间"),
        ("UpdatedAt", "DATETIME2", "", "更新时间"),
    ]),
    ("4.2 Songs 歌曲表", [
        ("Id", "INT", "PK", "自增主键"),
        ("Title", "NVARCHAR(200)", "", "歌曲名称"),
        ("Artist", "NVARCHAR(200)", "", "歌手"),
        ("Genre", "NVARCHAR(50)", "", "流派"),
        ("Language", "NVARCHAR(20)", "", "语种"),
        ("Duration", "INT", "", "时长（秒）"),
        ("FileSize", "BIGINT", "", "文件大小"),
        ("CoverUrl", "NVARCHAR(500)", "", "封面路径"),
        ("MediaUrl", "NVARCHAR(500)", "", "音频路径"),
        ("LrcUrl", "NVARCHAR(500)", "", "歌词路径"),
        ("OriginalFileName", "NVARCHAR(500)", "", "原始文件名"),
        ("PlayCount", "INT", "", "播放次数"),
        ("Status", "NVARCHAR(20)", "", "状态 (active/disabled)"),
        ("CreatedAt", "DATETIME2", "", "创建时间"),
        ("UpdatedAt", "DATETIME2", "", "更新时间"),
    ]),
    ("4.3 Rooms 虚拟房间表", [
        ("Id", "INT", "PK", "自增主键"),
        ("RoomCode", "NVARCHAR(10)", "UNIQUE", "6 位房间码"),
        ("Status", "NVARCHAR(20)", "", "状态 (active/closed)"),
        ("CreatedByUserId", "INT", "FK→Users", "创建者 ID"),
        ("CurrentUsers", "INT", "", "当前人数"),
        ("IdleCloseAt", "DATETIME2", "", "空闲关闭时间"),
        ("CreatedAt", "DATETIME2", "", "创建时间"),
        ("ClosedAt", "DATETIME2", "", "关闭时间"),
    ]),
    ("4.4 RoomUsers 房间用户关联表", [
        ("Id", "INT", "PK", "自增主键"),
        ("RoomId", "INT", "FK→Rooms", "房间 ID"),
        ("UserId", "INT", "FK→Users, UNIQUE", "用户 ID（一个用户同时只在一个房间）"),
        ("JoinedAt", "DATETIME2", "", "加入时间"),
    ]),
    ("4.5 RoomRequests 房间申请表", [
        ("Id", "INT", "PK", "自增主键"),
        ("UserId", "INT", "FK→Users", "申请人 ID"),
        ("Status", "NVARCHAR(20)", "", "状态 (pending/approved/rejected)"),
        ("RoomId", "INT", "FK→Rooms, NULL", "分配的房间 ID"),
        ("CreatedAt", "DATETIME2", "", "申请时间"),
        ("ProcessedAt", "DATETIME2", "", "处理时间"),
        ("ProcessedBy", "INT", "FK→Users, NULL", "处理人 ID"),
    ]),
    ("4.6 PlayQueue 播放队列表", [
        ("Id", "INT", "PK", "自增主键"),
        ("RoomId", "INT", "FK→Rooms", "房间 ID"),
        ("SongId", "INT", "FK→Songs", "歌曲 ID"),
        ("OrderedByUserId", "INT", "FK→Users", "点歌用户 ID"),
        ("SortOrder", "INT", "", "排序序号"),
        ("Status", "NVARCHAR(20)", "", "状态 (queued/removed)"),
        ("CreatedAt", "DATETIME2", "", "创建时间"),
    ]),
    ("4.7 Favorites 收藏表", [
        ("Id", "INT", "PK", "自增主键"),
        ("UserId", "INT", "FK→Users", "用户 ID"),
        ("SongId", "INT", "FK→Songs", "歌曲 ID"),
        ("CreatedAt", "DATETIME2", "", "收藏时间"),
    ]),
    ("4.8 Feedbacks 反馈表", [
        ("Id", "INT", "PK", "自增主键"),
        ("UserId", "INT", "FK→Users", "反馈用户 ID"),
        ("FeedbackType", "NVARCHAR(30)", "", "反馈类型"),
        ("SongName", "NVARCHAR(200)", "NULL", "歌曲名"),
        ("Artist", "NVARCHAR(200)", "NULL", "歌手"),
        ("Description", "NVARCHAR(1000)", "NULL", "描述"),
        ("Status", "NVARCHAR(20)", "", "状态 (pending/processed)"),
        ("CreatedAt", "DATETIME2", "", "创建时间"),
        ("ProcessedAt", "DATETIME2", "NULL", "处理时间"),
    ]),
    ("4.9 ChatMessages 聊天消息表", [
        ("Id", "INT", "PK", "自增主键"),
        ("RoomId", "INT", "FK→Rooms", "房间 ID"),
        ("UserId", "INT", "FK→Users", "发送者 ID"),
        ("Nickname", "NVARCHAR(50)", "", "发送者昵称"),
        ("Content", "NVARCHAR(200)", "", "消息内容"),
        ("CreatedAt", "DATETIME2", "", "发送时间"),
    ]),
    ("4.10 SystemSettings 系统设置表", [
        ("Id", "INT", "PK", "自增主键"),
        ("SettingKey", "NVARCHAR(100)", "UNIQUE", "设置键"),
        ("SettingValue", "NVARCHAR(MAX)", "", "设置值"),
        ("UpdatedAt", "DATETIME2", "", "更新时间"),
    ]),
    ("4.11 OperationLogs 操作日志表", [
        ("Id", "INT", "PK", "自增主键"),
        ("Username", "NVARCHAR(50)", "", "操作人"),
        ("OperationType", "NVARCHAR(50)", "", "操作类型"),
        ("ObjectType", "NVARCHAR(50)", "", "对象类型"),
        ("ObjectId", "NVARCHAR(50)", "NULL", "对象 ID"),
        ("Details", "NVARCHAR(MAX)", "NULL", "详情"),
        ("CreatedAt", "DATETIME2", "", "操作时间"),
    ]),
]

for title, fields in table_defs:
    doc.add_heading(title, level=2)
    t = doc.add_table(rows=len(fields) + 1, cols=4)
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(["字段名", "类型", "约束", "说明"]):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for i, (name, typ, constraint, desc) in enumerate(fields):
        t.rows[i + 1].cells[0].text = name
        t.rows[i + 1].cells[1].text = typ
        t.rows[i + 1].cells[2].text = constraint
        t.rows[i + 1].cells[3].text = desc
    doc.add_paragraph()

# ===== 5. 关系说明 =====
doc.add_heading("5. 关系说明", level=1)

doc.add_heading("5.1 核心业务关系", level=2)
biz_items = [
    "用户注册登录后，可浏览歌曲库、收藏歌曲、申请开房。",
    "管理员审批开房申请后，分配 6 位房间码，用户凭码加入房间。",
    "房间内用户可点歌（加入播放队列）、聊天、拖拽排序播放列表。",
    "播放队列按排序序号依次播放，播放后状态变为 removed。",
    "用户退出房间后，其点播的歌曲从队列中清除。",
]
for item in biz_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("5.2 数据一致性约束", level=2)
constraints = [
    "RoomUsers 表 UNIQUE(UserId) 约束：一个用户同时只能在一个房间。",
    "Favorites 表 UNIQUE(UserId, SongId) 约束：同一首歌只能收藏一次。",
    "Rooms 表 UNIQUE(RoomCode) 约束：房间码全局唯一。",
    "Users 表 UNIQUE(Username) 和 UNIQUE(Email) 约束：用户名和邮箱不能重复。",
]
for item in constraints:
    doc.add_paragraph(item, style="List Bullet")

# ===== 6. 状态值枚举 =====
doc.add_heading("6. 状态值枚举", level=1)

enum_table = doc.add_table(rows=9, cols=3)
enum_table.style = "Light Grid Accent 1"
for j, h in enumerate(["表.字段", "值", "含义"]):
    cell = enum_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

enums = [
    ("Users.Status", "active", "正常"),
    ("Users.Status", "disabled", "禁用"),
    ("Users.Role", "user", "普通用户"),
    ("Users.Role", "admin", "管理员"),
    ("Rooms.Status", "active", "进行中"),
    ("Rooms.Status", "closed", "已关闭"),
    ("PlayQueue.Status", "queued", "排队中"),
    ("PlayQueue.Status", "removed", "已移除"),
]
for i, (table_field, value, meaning) in enumerate(enums):
    enum_table.rows[i + 1].cells[0].text = table_field
    enum_table.rows[i + 1].cells[1].text = value
    enum_table.rows[i + 1].cells[2].text = meaning

# 保存
import os
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "docs", "ER-diagram.docx")
doc.save(output_path)
print(f"Word 文档已生成: {output_path}")
